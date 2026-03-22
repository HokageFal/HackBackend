from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import re
from typing import Dict, Any, List


class ReportTemplateEngine:
    """
    Движок для генерации отчетов по шаблонам.
    
    Поддерживаемые переменные в шаблоне:
    - {{client_name}} - ФИО клиента
    - {{test_title}} - Название теста
    - {{submitted_date}} - Дата прохождения
    - {{profile.field_name}} - Значение поля профиля
    - {{answer.question_id}} - Ответ на вопрос
    - {{metric.metric_name}} - Значение метрики
    - {{section.answers}} - Все ответы из секции
    """
    
    def __init__(self, template_definition: dict):
        self.template_definition = template_definition
    
    def generate_docx(
        self,
        test_title: str,
        client_name: str,
        submitted_at: str,
        profile_data: List[Dict],
        answers: List[Dict],
        calculated_metrics: Dict[str, Any],
        audience: str
    ) -> BytesIO:
        doc = Document()
        
        context = self._build_context(
            test_title=test_title,
            client_name=client_name,
            submitted_at=submitted_at,
            profile_data=profile_data,
            answers=answers,
            calculated_metrics=calculated_metrics,
            audience=audience
        )
        
        sections = self.template_definition.get("sections", [])
        
        for section in sections:
            self._render_section(doc, section, context)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    def _build_context(
        self,
        test_title: str,
        client_name: str,
        submitted_at: str,
        profile_data: List[Dict],
        answers: List[Dict],
        calculated_metrics: Dict[str, Any],
        audience: str
    ) -> Dict[str, Any]:
        profile_dict = {pd["field_name"]: pd["value"] for pd in profile_data}
        
        answers_dict = {}
        for ans in answers:
            answer_text = ""
            if ans.get("answer_value") is not None:
                if isinstance(ans["answer_value"], bool):
                    answer_text = "Да" if ans["answer_value"] else "Нет"
                else:
                    answer_text = str(ans["answer_value"])
            elif ans.get("selected_options"):
                answer_text = ", ".join([opt["option_text"] for opt in ans["selected_options"]])
            
            answers_dict[str(ans["question_id"])] = {
                "text": ans["question_text"],
                "answer": answer_text,
                "type": ans["question_type"]
            }
        
        return {
            "client_name": client_name,
            "test_title": test_title,
            "submitted_date": submitted_at,
            "audience": audience,
            "profile": profile_dict,
            "answers": answers_dict,
            "metrics": calculated_metrics,
            "all_answers": answers,
            "all_profile": profile_data
        }
    
    def _render_section(self, doc: Document, section: Dict, context: Dict):
        section_type = section.get("type", "text")
        
        if section_type == "heading":
            self._render_heading(doc, section, context)
        elif section_type == "text":
            self._render_text(doc, section, context)
        elif section_type == "table":
            self._render_table(doc, section, context)
        elif section_type == "profile_data":
            self._render_profile_data(doc, section, context)
        elif section_type == "answers_list":
            self._render_answers_list(doc, section, context)
        elif section_type == "metrics":
            self._render_metrics(doc, section, context)
        elif section_type == "page_break":
            doc.add_page_break()
    
    def _render_heading(self, doc: Document, section: Dict, context: Dict):
        content = section.get("content", "")
        level = section.get("level", 1)
        align = section.get("align", "left")
        
        content = self._replace_variables(content, context)
        
        heading = doc.add_heading(content, level=level)
        
        if align == "center":
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    def _render_text(self, doc: Document, section: Dict, context: Dict):
        content = section.get("content", "")
        bold = section.get("bold", False)
        italic = section.get("italic", False)
        
        content = self._replace_variables(content, context)
        
        p = doc.add_paragraph()
        run = p.add_run(content)
        run.bold = bold
        run.italic = italic
    
    def _render_table(self, doc: Document, section: Dict, context: Dict):
        rows = section.get("rows", [])
        
        if not rows:
            return
        
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Light Grid Accent 1'
        
        for i, row_data in enumerate(rows):
            for j, cell_data in enumerate(row_data):
                cell_content = self._replace_variables(str(cell_data), context)
                table.rows[i].cells[j].text = cell_content
    
    def _render_profile_data(self, doc: Document, section: Dict, context: Dict):
        doc.add_heading('Данные профиля', level=2)
        
        for pd in context["all_profile"]:
            p = doc.add_paragraph()
            p.add_run(f'{pd["field_name"]}: ').bold = True
            p.add_run(str(pd["value"]) if pd["value"] else "Не указано")
    
    def _render_answers_list(self, doc: Document, section: Dict, context: Dict):
        title = section.get("title", "Ответы на вопросы")
        show_question_numbers = section.get("show_numbers", True)
        
        doc.add_heading(title, level=2)
        
        for i, ans in enumerate(context["all_answers"], 1):
            question_p = doc.add_paragraph()
            
            if show_question_numbers:
                question_p.add_run(f'{i}. {ans["question_text"]}').bold = True
            else:
                question_p.add_run(ans["question_text"]).bold = True
            
            answer_p = doc.add_paragraph(style='List Bullet')
            
            if ans.get("answer_value") is not None:
                answer_text = str(ans["answer_value"])
                if isinstance(ans["answer_value"], bool):
                    answer_text = "Да" if ans["answer_value"] else "Нет"
                answer_p.add_run(f'Ответ: {answer_text}')
            
            if ans.get("selected_options"):
                options_text = ", ".join([opt["option_text"] for opt in ans["selected_options"]])
                answer_p.add_run(f'Выбрано: {options_text}')
            
            if not ans.get("answer_value") and not ans.get("selected_options"):
                answer_p.add_run('Ответ не предоставлен')
    
    def _render_metrics(self, doc: Document, section: Dict, context: Dict):
        title = section.get("title", "Метрики и результаты")
        
        if not context["metrics"]:
            return
        
        doc.add_heading(title, level=2)
        
        for name, data in context["metrics"].items():
            metric_p = doc.add_paragraph()
            metric_p.add_run(f'{name}: ').bold = True
            metric_p.add_run(f'{data["value"]}\n')
            if data.get("description"):
                metric_p.add_run(data["description"]).italic = True
    
    def _replace_variables(self, text: str, context: Dict) -> str:
        pattern = r'\{\{([^}]+)\}\}'
        
        def replacer(match):
            var_path = match.group(1).strip()
            return str(self._get_nested_value(context, var_path))
        
        return re.sub(pattern, replacer, text)
    
    def _get_nested_value(self, data: Dict, path: str) -> Any:
        keys = path.split('.')
        value = data
        
        for key in keys:
            if isinstance(value, dict):
                value = value.get(key, f"{{{{MISSING: {path}}}}}")
            else:
                return f"{{{{INVALID: {path}}}}}"
        
        return value if value is not None else ""


def get_default_template(audience: str) -> dict:
    """Возвращает дефолтный шаблон отчета"""
    
    if audience == "client":
        return {
            "title": "Отчет для клиента",
            "sections": [
                {
                    "type": "heading",
                    "content": "{{test_title}}",
                    "level": 0,
                    "align": "center"
                },
                {
                    "type": "text",
                    "content": "Клиент: {{client_name}}\nДата прохождения: {{submitted_date}}",
                    "bold": False
                },
                {
                    "type": "profile_data"
                },
                {
                    "type": "answers_list",
                    "title": "Ваши ответы",
                    "show_numbers": True
                },
                {
                    "type": "metrics",
                    "title": "Результаты"
                },
                {
                    "type": "text",
                    "content": "\nОтчет сгенерирован автоматически системой ПрофДНК",
                    "italic": True
                }
            ]
        }
    else:  # psychologist
        return {
            "title": "Отчет для психолога",
            "sections": [
                {
                    "type": "heading",
                    "content": "Отчет: {{test_title}}",
                    "level": 0,
                    "align": "center"
                },
                {
                    "type": "text",
                    "content": "Клиент: {{client_name}}\nДата прохождения: {{submitted_date}}\nТип отчета: Для психолога",
                    "bold": False
                },
                {
                    "type": "profile_data"
                },
                {
                    "type": "answers_list",
                    "title": "Ответы клиента",
                    "show_numbers": True
                },
                {
                    "type": "metrics",
                    "title": "Рассчитанные метрики"
                },
                {
                    "type": "text",
                    "content": "\nКонфиденциальный отчет для специалиста",
                    "italic": True
                }
            ]
        }
