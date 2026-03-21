from sqlalchemy.ext.asyncio import AsyncSession
from typing import Optional
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.logging_config import get_logger
from app.database.crud.test_crud import get_test_by_id
from app.database.crud.test_attempt_crud import get_attempts_by_test, get_attempt_by_id
from app.database.crud.user_answer_crud import get_answers_by_attempt
from app.database.crud.user_answer_option_crud import get_selected_options
from app.database.crud.test_attempt_profile_value_crud import get_profile_values_by_attempt
from app.database.crud.question_crud import get_question_by_id
from app.database.crud.question_option_crud import get_option_by_id
from app.database.crud.test_profile_field_crud import get_profile_fields_by_test
from app.database.crud.test_metric_crud import get_metrics_by_test
from app.database.crud.report_template_crud import get_report_templates_by_test

logger = get_logger(__name__)


class AttemptNotFound(Exception):
    def __init__(self, field: str, message: str):
        self.field = field
        self.message = message
        super().__init__(message)


class TestNotFound(Exception):
    def __init__(self, field: str, message: str):
        self.field = field
        self.message = message
        super().__init__(message)


class AccessDenied(Exception):
    def __init__(self, field: str, message: str):
        self.field = field
        self.message = message
        super().__init__(message)


def generate_docx_report(
    test_title: str,
    client_name: str,
    submitted_at: datetime,
    profile_data: list,
    answers: list,
    calculated_metrics: dict,
    audience: str
) -> BytesIO:
    doc = Document()
    
    title = doc.add_heading(f'Отчет: {test_title}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    info_section = doc.add_paragraph()
    info_section.add_run('Информация о прохождении\n').bold = True
    info_section.add_run(f'Клиент: {client_name}\n')
    info_section.add_run(f'Дата прохождения: {submitted_at.strftime("%d.%m.%Y %H:%M")}\n')
    info_section.add_run(f'Тип отчета: {"Для клиента" if audience == "client" else "Для психолога"}\n')
    
    doc.add_heading('Данные профиля', level=1)
    
    if profile_data:
        for pd in profile_data:
            p = doc.add_paragraph()
            p.add_run(f'{pd["field_name"]}: ').bold = True
            p.add_run(str(pd["value"]) if pd["value"] else "Не указано")
    else:
        doc.add_paragraph('Данные профиля не заполнены')
    
    doc.add_heading('Ответы на вопросы', level=1)
    
    for i, ans in enumerate(answers, 1):
        question_p = doc.add_paragraph()
        question_p.add_run(f'{i}. {ans["question_text"]}').bold = True
        
        answer_p = doc.add_paragraph(style='List Bullet')
        
        if ans["answer_value"] is not None:
            answer_text = str(ans["answer_value"])
            if isinstance(ans["answer_value"], bool):
                answer_text = "Да" if ans["answer_value"] else "Нет"
            answer_p.add_run(f'Ответ: {answer_text}')
        
        if ans["selected_options"]:
            options_text = ", ".join([opt["option_text"] for opt in ans["selected_options"]])
            answer_p.add_run(f'Выбрано: {options_text}')
        
        if not ans["answer_value"] and not ans["selected_options"]:
            answer_p.add_run('Ответ не предоставлен')
    
    if calculated_metrics:
        doc.add_heading('Метрики и результаты', level=1)
        
        for name, data in calculated_metrics.items():
            metric_p = doc.add_paragraph()
            metric_p.add_run(f'{name}: ').bold = True
            metric_p.add_run(f'{data["value"]}\n')
            if data.get("description"):
                metric_p.add_run(data["description"]).italic = True
    
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run('\nОтчет сгенерирован автоматически системой ПрофДНК')
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


async def get_test_attempts_service(
    session: AsyncSession,
    test_id: int,
    psychologist_id: int,
    skip: int = 0,
    limit: int = 20
) -> dict:
    logger.info(
        "Starting test attempts retrieval",
        operation="get_test_attempts_service",
        test_id=test_id,
        psychologist_id=psychologist_id,
        skip=skip,
        limit=limit
    )
    
    try:
        test = await get_test_by_id(session, test_id)
        
        if not test:
            logger.warning(
                "Test not found",
                operation="get_test_attempts_service",
                test_id=test_id
            )
            raise TestNotFound("test_id", f"Тест с ID {test_id} не найден")
        
        if test.psychologist_id != psychologist_id:
            logger.warning(
                "Access denied to test attempts",
                operation="get_test_attempts_service",
                test_id=test_id,
                psychologist_id=psychologist_id,
                test_owner_id=test.psychologist_id
            )
            raise AccessDenied("test_id", "У вас нет доступа к результатам этого теста")
        
        attempts, total = await get_attempts_by_test(session, test_id, skip, limit)
        
        attempts_list = []
        for attempt in attempts:
            attempts_list.append({
                "id": attempt.id,
                "client_name": attempt.client_name,
                "started_at": attempt.started_at.isoformat(),
                "submitted_at": attempt.submitted_at.isoformat() if attempt.submitted_at else None,
                "status": "completed" if attempt.submitted_at else "in_progress"
            })
        
        total_pages = (total + limit - 1) // limit
        
        result = {
            "test_id": test_id,
            "test_title": test.title,
            "attempts": attempts_list,
            "pagination": {
                "total": total,
                "page": (skip // limit) + 1,
                "limit": limit,
                "total_pages": total_pages,
                "has_next": skip + limit < total,
                "has_prev": skip > 0
            }
        }
        
        logger.info(
            "Test attempts retrieved successfully",
            operation="get_test_attempts_service",
            test_id=test_id,
            count=len(attempts_list),
            total=total
        )
        
        return result
        
    except (TestNotFound, AccessDenied):
        raise
    except Exception as e:
        logger.error(
            "Unexpected error during test attempts retrieval",
            operation="get_test_attempts_service",
            test_id=test_id,
            error_type=type(e).__name__,
            error_message=str(e),
            exc_info=True
        )
        raise


async def get_attempt_details_service(
    session: AsyncSession,
    attempt_id: int,
    psychologist_id: int
) -> dict:
    logger.info(
        "Starting attempt details retrieval",
        operation="get_attempt_details_service",
        attempt_id=attempt_id,
        psychologist_id=psychologist_id
    )
    
    try:
        attempt = await get_attempt_by_id(session, attempt_id)
        
        if not attempt:
            logger.warning(
                "Attempt not found",
                operation="get_attempt_details_service",
                attempt_id=attempt_id
            )
            raise AttemptNotFound("attempt_id", f"Попытка с ID {attempt_id} не найдена")
        
        test = await get_test_by_id(session, attempt.test_id)
        
        if test.psychologist_id != psychologist_id:
            logger.warning(
                "Access denied to attempt details",
                operation="get_attempt_details_service",
                attempt_id=attempt_id,
                psychologist_id=psychologist_id,
                test_owner_id=test.psychologist_id
            )
            raise AccessDenied("attempt_id", "У вас нет доступа к результатам этой попытки")
        
        profile_values = await get_profile_values_by_attempt(session, attempt_id)
        profile_fields = await get_profile_fields_by_test(session, attempt.test_id)
        
        profile_data = []
        for pv in profile_values:
            field = next((pf for pf in profile_fields if pf.id == pv.profile_field_id), None)
            if field:
                value = pv.text_value or pv.number_value or pv.date_value
                profile_data.append({
                    "field_id": pv.profile_field_id,
                    "field_name": field.label,
                    "field_type": field.type.value,
                    "value": str(value) if value else None
                })
        
        user_answers = await get_answers_by_attempt(session, attempt_id)
        
        answers_data = []
        for answer in user_answers:
            question = await get_question_by_id(session, answer.question_id)
            
            if not question:
                continue
            
            answer_value = None
            selected_options = []
            
            if answer.text_answer:
                answer_value = answer.text_answer
            elif answer.boolean_answer is not None:
                answer_value = answer.boolean_answer
            elif answer.number_answer is not None:
                answer_value = float(answer.number_answer)
            elif answer.datetime_answer:
                answer_value = answer.datetime_answer.isoformat()
            
            answer_options = await get_selected_options(session, answer.id)
            for ao in answer_options:
                option = await get_option_by_id(session, ao.option_id)
                if option:
                    selected_options.append({
                        "option_id": option.id,
                        "option_text": option.text,
                        "option_value": option.position
                    })
            
            answers_data.append({
                "question_id": question.id,
                "question_text": question.text,
                "question_type": question.type.value,
                "answer_value": answer_value,
                "selected_options": selected_options if selected_options else None
            })
        
        result = {
            "attempt_id": attempt.id,
            "test_id": attempt.test_id,
            "test_title": test.title,
            "client_name": attempt.client_name,
            "started_at": attempt.started_at.isoformat(),
            "submitted_at": attempt.submitted_at.isoformat() if attempt.submitted_at else None,
            "status": "completed" if attempt.submitted_at else "in_progress",
            "profile_data": profile_data,
            "answers": answers_data
        }
        
        logger.info(
            "Attempt details retrieved successfully",
            operation="get_attempt_details_service",
            attempt_id=attempt_id,
            answers_count=len(answers_data)
        )
        
        return result
        
    except (AttemptNotFound, AccessDenied):
        raise
    except Exception as e:
        logger.error(
            "Unexpected error during attempt details retrieval",
            operation="get_attempt_details_service",
            attempt_id=attempt_id,
            error_type=type(e).__name__,
            error_message=str(e),
            exc_info=True
        )
        raise


async def generate_report_service(
    session: AsyncSession,
    attempt_id: int,
    psychologist_id: int,
    audience: str,
    format: str = "html"
):
    logger.info(
        "Starting report generation",
        operation="generate_report_service",
        attempt_id=attempt_id,
        audience=audience,
        format=format
    )
    
    try:
        attempt = await get_attempt_by_id(session, attempt_id)
        
        if not attempt:
            raise AttemptNotFound("attempt_id", f"Попытка с ID {attempt_id} не найдена")
        
        if not attempt.submitted_at:
            raise AccessDenied("attempt_id", "Отчет можно сформировать только для завершенных тестов")
        
        test = await get_test_by_id(session, attempt.test_id)
        
        if test.psychologist_id != psychologist_id:
            raise AccessDenied("attempt_id", "У вас нет доступа к результатам этой попытки")
        
        attempt_details = await get_attempt_details_service(session, attempt_id, psychologist_id)
        
        metrics = await get_metrics_by_test(session, attempt.test_id)
        calculated_metrics = {}
        for metric in metrics:
            calculated_metrics[metric.metric_name] = {
                "value": 0,
                "formula": metric.formula,
                "description": metric.description
            }
        
        templates = await get_report_templates_by_test(session, attempt.test_id)
        template = next((t for t in templates if t.audience.value == audience), None)
        
        if not template:
            logger.info(
                "Report template not found, using default",
                operation="generate_report_service",
                attempt_id=attempt_id,
                audience=audience
            )
        
        if format == "docx":
            docx_buffer = generate_docx_report(
                test_title=test.title,
                client_name=attempt.client_name,
                submitted_at=attempt.submitted_at,
                profile_data=attempt_details["profile_data"],
                answers=attempt_details["answers"],
                calculated_metrics=calculated_metrics,
                audience=audience
            )
            
            logger.info(
                "DOCX report generated successfully",
                operation="generate_report_service",
                attempt_id=attempt_id,
                audience=audience
            )
            
            return docx_buffer
        
        report_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Отчет - {test.title}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                h1 {{ color: #333; }}
                h2 {{ color: #666; margin-top: 30px; }}
                .info {{ background: #f5f5f5; padding: 15px; border-radius: 5px; margin: 20px 0; }}
                .answer {{ margin: 15px 0; padding: 10px; border-left: 3px solid #4CAF50; }}
                .question {{ font-weight: bold; color: #333; }}
                .value {{ color: #666; margin-top: 5px; }}
                .metric {{ background: #e3f2fd; padding: 10px; margin: 10px 0; border-radius: 5px; }}
            </style>
        </head>
        <body>
            <h1>Отчет: {test.title}</h1>
            
            <div class="info">
                <p><strong>Клиент:</strong> {attempt.client_name}</p>
                <p><strong>Дата прохождения:</strong> {attempt.submitted_at.strftime('%d.%m.%Y %H:%M') if attempt.submitted_at else 'Не завершен'}</p>
            </div>
            
            <h2>Данные профиля</h2>
            {''.join([f'<p><strong>{pd["field_name"]}:</strong> {pd["value"]}</p>' for pd in attempt_details["profile_data"]])}
            
            <h2>Ответы на вопросы</h2>
            {''.join([f'''
            <div class="answer">
                <div class="question">{i+1}. {ans["question_text"]}</div>
                <div class="value">
                    {ans["answer_value"] if ans["answer_value"] is not None else ""}
                    {", ".join([opt["option_text"] for opt in ans["selected_options"]]) if ans["selected_options"] else ""}
                </div>
            </div>
            ''' for i, ans in enumerate(attempt_details["answers"])])}
            
            <h2>Метрики</h2>
            {''.join([f'<div class="metric"><strong>{name}:</strong> {data["value"]} <br><small>{data["description"]}</small></div>' for name, data in calculated_metrics.items()])}
            
            <p style="margin-top: 40px; color: #999; font-size: 12px;">
                Отчет сгенерирован автоматически системой ПрофДНК
            </p>
        </body>
        </html>
        """
        
        result = {
            "attempt_id": attempt_id,
            "test_title": test.title,
            "client_name": attempt.client_name,
            "audience": audience,
            "format": format,
            "generated_at": datetime.utcnow().isoformat(),
            "content": report_html
        }
        
        logger.info(
            "HTML report generated successfully",
            operation="generate_report_service",
            attempt_id=attempt_id,
            audience=audience
        )
        
        return result
        
    except (AttemptNotFound, AccessDenied):
        raise
    except Exception as e:
        logger.error(
            "Unexpected error during report generation",
            operation="generate_report_service",
            attempt_id=attempt_id,
            error_type=type(e).__name__,
            error_message=str(e),
            exc_info=True
        )
        raise
